"""E2E REAL de ejecución de skills en el loop: bash corre los scripts bundled de
dos skills de oficina y genera artefactos `.docx`/`.xlsx` reales en disco.

A diferencia de `test_capability_registration_e2e` (que prueba el FLUJO controlado por el
runtime: registrar → encontrar → invocar Skill → habilitar tools → exponer base_dir), aquí
se cierra el último eslabón: **bash ejecuta de verdad los scripts** de las skills y el test
verifica los artefactos reabriéndolos con `python-docx`/`openpyxl`.

Lo único guionizado es la decisión del modelo (qué tool llamar): el caller scripted lee el
`base_dir` que el loop SURFACE en el tool-result de `Skill` y construye desde ahí el comando
bash — igual que haría el modelo real leyendo "Base directory for this skill: <dir>". La
ejecución del subproceso y la generación del fichero son reales y se verifican por contenido.

Cadena ejercitada por skill (×2, docx y xlsx):
  turno N   → Skill(officeX)  → habilita `bash` (allowed-tools) + surface base_dir
  turno N+1 → bash <interprete> <base_dir>/scripts/make_X.py <out> <token>  → crea el fichero
  …
  turno fin → texto

Las skills se registran como DIRECTORIO en disco (`<root>/<skill>/SKILL.md` + `scripts/`),
que es lo que fija `base_dir` al `skillRoot` real — sin él, el modelo no localizaría los
scripts. `bash` solo queda permitido porque la skill lo declara en `allowed-tools`.

Deps de las skills (`python-docx`, `openpyxl`): grupo `skills-e2e` del pyproject — son las
`deps` que el integrador provee para estas skills, NO dependencias del runtime. El test se
omite si no están instaladas.
"""
from __future__ import annotations

import sys
import uuid
from pathlib import Path

import pytest

from agentic_runtime.contracts.runtime import RuntimeTask
from agentic_runtime.events import DoneEvent, TokenEvent, ToolCallEvent
from agentic_runtime.factory import CapabilitiesConfig, RuntimeConfig, StorageConfig, create_runtime

# Gating: sin las deps de las skills no hay nada real que ejecutar.
docx = pytest.importorskip("docx", reason="falta python-docx (grupo skills-e2e)")
openpyxl = pytest.importorskip("openpyxl", reason="falta openpyxl (grupo skills-e2e)")


_MAKE_DOCX = """\
import sys
from docx import Document

out, token = sys.argv[1], sys.argv[2]
doc = Document()
doc.add_heading("Informe trimestral", level=1)
doc.add_paragraph(f"token:{token}")
doc.save(out)
print(f"docx escrito en {out}")
"""

_MAKE_XLSX = """\
import sys
from openpyxl import Workbook

out, token = sys.argv[1], sys.argv[2]
wb = Workbook()
ws = wb.active
ws["A1"] = "token"
ws["B1"] = token
wb.save(out)
print(f"xlsx escrito en {out}")
"""


def _skill_md(name: str, desc: str, script: str) -> str:
    return (
        f"---\n"
        f"name: {name}\n"
        f"description: {desc}\n"
        f"allowed-tools:\n"
        f"  - bash\n"
        f"---\n"
        f"Para generar el documento, ejecuta con bash el script `scripts/{script}` que vive\n"
        f"en el base directory de esta skill, pasándole la ruta de salida y un token.\n"
    )


def _write_skill_dir(root: Path, name: str, desc: str, script: str, body: str) -> None:
    skill_dir = root / name
    (skill_dir / "scripts").mkdir(parents=True)
    (skill_dir / "SKILL.md").write_text(_skill_md(name, desc, script), encoding="utf-8")
    (skill_dir / "scripts" / script).write_text(body, encoding="utf-8")


def _last_base_dir(messages: list) -> str | None:
    """Extrae el `base_dir` más reciente surface por el loop en un tool-result de Skill.

    Espejo de lo que haría el modelo real al leer "Base directory for this skill: <dir>".
    """
    base: str | None = None
    for m in messages:
        content = str(m.get("content", "") if isinstance(m, dict) else getattr(m, "content", ""))
        for line in content.splitlines():
            if line.startswith("Base directory for this skill:"):
                base = line.split(":", 1)[1].strip()
    return base


class _OfficeCaller:
    """Conduce el loop: por cada skill, Skill(officeX) → bash(script) → (siguiente) → fin.

    En el turno de bash NO inventa la ruta: parsea el `base_dir` que el loop surface en el
    tool-result de `Skill` y compone `<interprete> <base_dir>/scripts/make_X.py <out> <token>`.
    Usa `sys.executable` para garantizar el intérprete con las deps de las skills.
    """

    def __init__(self, plan: list[tuple[str, str, Path, str]]) -> None:
        # plan: [(skill_name, script_rel, out_path, token), …]
        self._plan = plan
        self.turns: list[list[str]] = []
        self.bash_commands: list[str] = []
        self.last_messages: list = []

    async def complete(self, messages, tools, *, stop=None, model_id=""):
        self.turns.append([t["name"] for t in tools])
        self.last_messages = list(messages)
        n = len(self.turns)

        # Cada skill consume 2 turnos: (2k+1) Skill, (2k+2) bash. Resto: fin.
        async def gen():
            if n % 2 == 1 and (n // 2) < len(self._plan):
                skill_name, _script, _out, _token = self._plan[n // 2]
                yield ToolCallEvent(
                    tool_name="Skill", tool_input={"command": skill_name}, call_id=f"s{n}"
                )
                yield DoneEvent(stop_reason="tool_calls")
            elif n % 2 == 0 and (n // 2 - 1) < len(self._plan):
                _skill_name, script, out, token = self._plan[n // 2 - 1]
                base = _last_base_dir(self.last_messages)
                assert base, "el loop no surface el base_dir de la skill en el tool-result"
                cmd = f'"{sys.executable}" "{Path(base) / "scripts" / script}" "{out}" {token}'
                self.bash_commands.append(cmd)
                yield ToolCallEvent(
                    tool_name="bash", tool_input={"command": cmd}, call_id=f"b{n}"
                )
                yield DoneEvent(stop_reason="tool_calls")
            else:
                yield TokenEvent(content="documentos generados")
                yield DoneEvent(stop_reason="stop")

        return gen()


async def test_office_skills_run_scripts_in_loop_and_produce_real_files(tmp_path):
    skills_root = tmp_path / "skills_src"
    out_dir = tmp_path / "out"
    out_dir.mkdir(parents=True)

    _write_skill_dir(skills_root, "officedocx", "crea documentos .docx", "make_docx.py", _MAKE_DOCX)
    _write_skill_dir(skills_root, "officexlsx", "crea hojas .xlsx", "make_xlsx.py", _MAKE_XLSX)

    docx_token = f"DOCX-{uuid.uuid4().hex[:8].upper()}"
    xlsx_token = f"XLSX-{uuid.uuid4().hex[:8].upper()}"
    docx_out = out_dir / "informe.docx"
    xlsx_out = out_dir / "datos.xlsx"

    caller = _OfficeCaller([
        ("officedocx", "make_docx.py", docx_out, docx_token),
        ("officexlsx", "make_xlsx.py", xlsx_out, xlsx_token),
    ])

    runtime = create_runtime(config=RuntimeConfig(
        storage=StorageConfig(backend="filesystem", root=tmp_path),
        model_caller=caller,
        # Skills registradas como DIRECTORIO en disco → base_dir = skillRoot real.
        capabilities=CapabilitiesConfig(skill_dirs=[skills_root]),
    ))
    await runtime.startup()
    try:
        task_id = await runtime.dispatch(RuntimeTask(prompt="genera los documentos", description="office-e2e"))
        rec = runtime._task_registry.get(task_id)
        await rec.asyncio_task
    finally:
        await runtime.shutdown()

    # --- anuncio homologado al canónico: bash (requires_permission) se anuncia desde el
    #     inicio; su gate vive en ejecución (dispatcher + hook), no en la visibilidad ---
    assert "Skill" in caller.turns[0]
    assert "bash" in caller.turns[0]              # gated pero visible: visibilidad ⟂ permiso
    assert "bash" in caller.turns[1]              # sigue anunciado tras Skill(officedocx)
    assert "bash" in caller.turns[3]              # idem tras Skill(officexlsx)

    # --- bash ejecutó de verdad los scripts (sin error) ---
    tool_results = [e for e in rec.events if e["type"] == "tool_result"]
    bash_results = [e for e in tool_results if e["call_id"].startswith("b")]
    assert len(bash_results) == 2
    assert all(not e["is_error"] for e in bash_results), bash_results

    # --- artefactos REALES en disco, verificados reabriéndolos por contenido ---
    assert docx_out.exists() and xlsx_out.exists()

    doc = docx.Document(str(docx_out))
    doc_text = "\n".join(p.text for p in doc.paragraphs)
    assert f"token:{docx_token}" in doc_text     # token único: solo el script pudo escribirlo

    wb = openpyxl.load_workbook(str(xlsx_out))
    ws = wb.active
    assert ws["A1"].value == "token"
    assert ws["B1"].value == xlsx_token
